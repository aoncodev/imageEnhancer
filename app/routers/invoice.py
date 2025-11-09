from fastapi import APIRouter, HTTPException, Request
from fastapi.responses import JSONResponse
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from app.services.s3_service import S3Service
import os
import tempfile
import re
import requests
from io import BytesIO
from PIL import Image
import json
from datetime import datetime

router = APIRouter()
s3 = S3Service()

@router.post("/generate-docx")
async def generate_docx(request: Request):
    """Generate invoice with bolded text, formatted numbers, and logo/seal images."""
    try:
        data = await request.json()
        print("\n📦 Incoming JSON data:")
        print(json.dumps(data, indent=2, ensure_ascii=False))

        if not isinstance(data, dict) or not data:
            raise HTTPException(status_code=400, detail="Request body must be a non-empty JSON object")

        return await generate_docx_from_data(data)

    except Exception as e:
        print(f"❌ ERROR: {e}")
        raise HTTPException(status_code=500, detail=str(e))


# -------------------------------------------------------------
# 🧠 Core logic — formatting, bold text, images, and centering
# -------------------------------------------------------------
async def generate_docx_from_data(data: dict):
    try:
        # ---------------------------------------------------------
        # 📄 Download template from S3
        # ---------------------------------------------------------
        template_url = data.get("file_url") or data.get("template_url")

        if not template_url:
            raise HTTPException(
                status_code=400,
                detail="file_url (or legacy template_url) is required in request body",
            )

        try:
            template_resp = requests.get(template_url, timeout=30)
            template_resp.raise_for_status()
            template_data = BytesIO(template_resp.content)
            doc = Document(template_data)
            print(f"📄 Loaded template from S3: {template_url}")
        except Exception as e:
            raise HTTPException(status_code=400, detail=f"Failed to download template from S3: {str(e)}")

        # ---------------------------------------------------------
        # 🖼️ Download images (logo + seal)
        # ---------------------------------------------------------
        def download_image(url):
            resp = requests.get(url, timeout=10)
            resp.raise_for_status()
            return BytesIO(resp.content)

        logo_data, seal_data = None, None
        logo_height_in = None

        if data.get("logo_image"):
            try:
                logo_data = download_image(data["logo_image"])
                with Image.open(logo_data) as img:
                    w, h = img.size
                    logo_height_in = (h / w) * 1.3
                logo_data.seek(0)
                print(f"✅ Logo downloaded (height: {logo_height_in:.2f}in)")
            except Exception as e:
                print(f"⚠️ Failed to download logo: {e}")

        if data.get("seal_image"):
            try:
                seal_data = download_image(data["seal_image"])
                print("✅ Seal downloaded")
            except Exception as e:
                print(f"⚠️ Failed to download seal: {e}")

        # ---------------------------------------------------------
        # 🔢 Format numeric values with commas
        # ---------------------------------------------------------
        for key in ["unit_price", "weight", "volume"]:
            if key in data and data[key]:
                try:
                    val_str = str(data[key]).replace(",", "").strip()
                    if val_str:
                        val = float(val_str)
                        if val.is_integer():
                            data[key] = f"{int(val):,}"
                        else:
                            data[key] = f"{val:,.2f}"
                        print(f"🔢 Formatted {key}: {data[key]}")
                except (ValueError, TypeError) as e:
                    print(f"⚠️ Could not format {key}: {e}")

        # ---------------------------------------------------------
        # 📅 Format date values as text (e.g., 20 November 2025)
        # ---------------------------------------------------------
        def format_date_text(value: str, include_year: bool = True) -> str:
            try:
                value_clean = str(value).strip()
                if not value_clean:
                    return value
                dt = datetime.strptime(value_clean, "%Y-%m-%d")
                return dt.strftime("%d %B %Y") if include_year else dt.strftime("%d %B")
            except (ValueError, TypeError):
                return value

        if "invoice_date" in data and data["invoice_date"]:
            original_invoice = data["invoice_date"]
            data["invoice_date"] = format_date_text(data["invoice_date"], include_year=True)
            if data["invoice_date"] != original_invoice:
                print(f"📅 Formatted invoice_date: {data['invoice_date']}")

        if "sailing_date" in data and data["sailing_date"]:
            original_sailing = data["sailing_date"]
            data["sailing_date"] = format_date_text(data["sailing_date"], include_year=False)
            if data["sailing_date"] != original_sailing:
                print(f"📅 Formatted sailing_date: {data['sailing_date']}")

        # ---------------------------------------------------------
        # 🌍 Swap destination and destination_country values
        # ---------------------------------------------------------
        if "destination" in data or "destination_country" in data:
            destination_value = data.get("destination")
            destination_country_value = data.get("destination_country")
            data["destination"] = destination_country_value
            data["destination_country"] = destination_value

        # ---------------------------------------------------------
        # ✍️ Replace text placeholders with bold values
        # Supports {key} format
        # Preserves paragraph spacing and font - only makes text bold
        # Special handling for port_loading: includes port_name + port_loading on separate lines
        # ---------------------------------------------------------
        def replace_text_in_paragraph(paragraph):
            full_text = "".join(run.text for run in paragraph.runs)
            if not re.search(r"\{\s*\w+\s*\}", full_text):
                return

            new_text = full_text
            
            # Special handling for port_loading: combine port_name and port_loading
            if re.search(r"\{\s*port_loading\s*\}", new_text, re.IGNORECASE):
                port_name = data.get("port_name", "")
                port_loading = data.get("port_loading", "")
                port_value = ""
                if port_name and port_loading:
                    port_value = f"{port_name}\n{port_loading}"
                elif port_name:
                    port_value = port_name
                elif port_loading:
                    port_value = port_loading
                
                if port_value:
                    pattern = re.compile(r"\{\s*port_loading\s*\}", re.IGNORECASE)
                    new_text = pattern.sub(port_value, new_text)
            
            # Replace other placeholders
            for key, value in data.items():
                if key in ["logo_image", "seal_image", "template_url", "file_url", "port_loading"]:
                    continue
                pattern = re.compile(rf"\{{\s*{re.escape(key)}\s*\}}", re.IGNORECASE)
                new_text = pattern.sub(str(value), new_text)

            if new_text != full_text:
                # Preserve paragraph formatting (spacing, alignment, font, etc.)
                # Get the original font from the first run to preserve it
                original_font_name = None
                original_font_size = None
                if paragraph.runs:
                    original_run = paragraph.runs[0]
                    try:
                        if original_run.font.name:
                            original_font_name = original_run.font.name
                    except:
                        pass
                    try:
                        if original_run.font.size:
                            original_font_size = original_run.font.size
                    except:
                        pass
                
                # Only replace the text content and make it bold
                for run in paragraph.runs[:]:
                    paragraph._element.remove(run._element)
                
                # Split by newlines to preserve line breaks
                lines = new_text.split('\n')
                for i, line in enumerate(lines):
                    # Add the line text
                    run = paragraph.add_run(line)
                    run.bold = True
                    # Preserve original font if it existed
                    if original_font_name:
                        try:
                            run.font.name = original_font_name
                        except:
                            pass
                    if original_font_size:
                        try:
                            run.font.size = original_font_size
                        except:
                            pass
                    # Add line break after each line except the last one
                    if i < len(lines) - 1:
                        paragraph.add_run().add_break()

        def replace_text_in_cell(cell):
            for paragraph in cell.paragraphs:
                replace_text_in_paragraph(paragraph)

        for paragraph in doc.paragraphs:
            replace_text_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    replace_text_in_cell(cell)

        # ---------------------------------------------------------
        # 🖼️ Insert images
        # Supports {key} format
        # ---------------------------------------------------------
        def insert_images_in_paragraph(paragraph):
            full_text = "".join(run.text for run in paragraph.runs)
            has_logo = re.search(r"\{\s*logo_image\s*\}", full_text, re.IGNORECASE)
            has_seal = re.search(r"\{\s*seal_image\s*\}", full_text, re.IGNORECASE)
            if not (has_logo or has_seal):
                return

            parts = re.split(r"(\{\s*logo_image\s*\}|\{\s*seal_image\s*\})", full_text, flags=re.IGNORECASE)
            for run in paragraph.runs[:]:
                paragraph._element.remove(run._element)

            for part in parts:
                if re.match(r"\{\s*logo_image\s*\}", part, re.IGNORECASE) and logo_data:
                    logo_data.seek(0)
                    run = paragraph.add_run()
                    run.add_picture(logo_data, width=Inches(1.3))
                elif re.match(r"\{\s*seal_image\s*\}", part, re.IGNORECASE) and seal_data:
                    seal_data.seek(0)
                    run = paragraph.add_run()
                    run.add_picture(seal_data, width=Inches(1.2))
                elif part.strip():
                    run = paragraph.add_run(part)
                    run.bold = True

        def insert_images_in_cell(cell):
            for paragraph in cell.paragraphs:
                insert_images_in_paragraph(paragraph)

        for paragraph in doc.paragraphs:
            insert_images_in_paragraph(paragraph)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    insert_images_in_cell(cell)

        # ---------------------------------------------------------
        # 💪 Make all text bold
        # ---------------------------------------------------------
        for paragraph in doc.paragraphs:
            for run in paragraph.runs:
                run.bold = True
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for paragraph in cell.paragraphs:
                        for run in paragraph.runs:
                            run.bold = True

        # ---------------------------------------------------------
        # 🧭 Center the QTY / UNIT PRICE / WEIGHT / VOLUME block
        # ---------------------------------------------------------
        def center_table_headers(doc):
            target_headers = ["QTY", "UNIT PRICE", "WEIGHT", "VOLUME"]
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            text = paragraph.text.strip().upper()
                            if any(header in text for header in target_headers):
                                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                for run in paragraph.runs:
                                    run.bold = True

        center_table_headers(doc)

        # ---------------------------------------------------------
        # 💾 Save and upload
        # ---------------------------------------------------------
        with tempfile.NamedTemporaryFile(delete=False, suffix=".docx") as tmp:
            temp_path = tmp.name
            doc.save(temp_path)

        print(f"💾 Document saved: {temp_path}")
        key = s3.generate_key("invoice_generated.docx", folder="docx-temp")
        with open(temp_path, "rb") as f:
            s3_url = s3.upload_bytes(
                data=f.read(),
                key=key,
                content_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
            )

        os.unlink(temp_path)
        print(f"✅ Uploaded to S3: {s3_url}")

        return JSONResponse({
            "message": "Invoice generated successfully",
            "download_url": s3_url
        })

    except Exception as e:
        print(f"❌ ERROR during generation: {e}")
        import traceback
        traceback.print_exc()
        raise HTTPException(status_code=500, detail=str(e))